from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_red_lines(file_path):
    doc = Document(file_path)
    
    # Add a page break to ensure it's on a new page or just a clear separation
    doc.add_page_break()
    
    # Heading
    heading = doc.add_heading('⚠️ CRITICAL RED LINES (绝不能踩的坑)', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(255, 0, 0) # Red
        run.bold = True

    red_lines = [
        ("🚫 No Oral Agreements", "Absolutely forbid oral promises with PLN officials, land owners, or contractors. Everything must be in written, stamped, and notarized contracts."),
        ("💰 No 100% Upfront Payments", "Forbid paying contractors 100% upfront. Implement a milestone-based payment system (e.g., 30% Deposit -> 40% Installation -> 30% SLO Certification)."),
        ("🇮🇩 TKDN Compliance", "Warn against importing fully assembled units for government projects without a local assembly partner, as this leads to immediate disqualification."),
        ("🌧️ Weather-proofing", "Explicitly warn against ignoring the 'Rainy Season' (Musim Hujan) risks—all outdoor equipment must be elevated 15cm+ from the ground to avoid flooding."),
        ("🔌 PLN Dependency", "Warn against relying on a single Fixer; establish a backup communication channel with the local PLN office.")
    ]

    for title, desc in red_lines:
        p = doc.add_paragraph()
        # Title run
        run_title = p.add_run(f"{title}: ")
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(255, 0, 0)
        run_title.font.size = Pt(12)
        
        # Description run
        run_desc = p.add_run(desc)
        run_desc.font.size = Pt(11)

    doc.save(file_path)

if __name__ == "__main__":
    path = "/Users/openclaw/.openclaw/workspace-shrimp-commander/Indonesia_EV_Charge_Pro_Operational_Manual_Final.docx"
    add_red_lines(path)
    print("Successfully updated document.")
