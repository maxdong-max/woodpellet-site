import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name, size=12):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_font(run, 'SimHei', 16)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 2:
        set_font(run, 'SimHei', 14)
    else:
        set_font(run, 'SimSun', 12)
    return p

def add_body_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 'SimSun', 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_table(doc, data):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    for r_idx, row_data in enumerate(data):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            run = cell.paragraphs[0].add_run(cell_text)
            set_font(run, 'SimSun', 10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def process_file(input_text, output_path):
    doc = Document()
    lines = input_text.split('\n')
    
    # Title
    first_line = lines[0].strip('# ').strip()
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(first_line)
    set_font(title_run, 'SimHei', 20)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    i = 1
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith('## '):
            add_heading(doc, line.replace('## ', '').strip(), 1)
        elif line.startswith('### '):
            add_heading(doc, line.replace('### ', '').strip(), 2)
        elif line.startswith('|'):
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].split('|') if cell.strip()]
                if not (row and all('---' in cell for cell in row if cell)):
                    table_data.append(row)
                i += 1
            add_table(doc, table_data)
            i -= 1
        else:
            if any(line.startswith(p) for p in ['1.', '(1)', '2.', '(2)']):
                add_heading(doc, line, 3 if line.startswith('1.') or line.startswith('2.') else 4)
            else:
                add_body_text(doc, line)
        i += 1
    doc.save(output_path)

report_text = """# 关于印尼电动汽车（EV）补能市场的量化研究报告

**版本：** 1.0（最终版）
**状态：** 最终同步版
**旨在：** 提供纯量化数据支撑，驱动战略决策与运营执行。

---

## 一、 市场量化分析

### （一） 补能基础设施增速与缺口
1. **公共充电站（SPKLU）增长率**
   2023年为1,081台，至2024年增至3,233台，年增长率近300%。
2. **2030年需求预测**
   预计需求量将达到3.2万台。
3. **市场价值量化**
   预计市场规模将从2024年的7.34亿美元增长至2030年的29.2亿美元，复合年增长率（CAGR）为22.85%。
4. **补能缺口分析**
   当前市场以中低速充电桩为主，120kW及以上高功率快充设备极度匮乏。大雅加达地区（Jabodetabek）顶级流量场站存在严重的“准入真空期”。

### （二） 竞争格局与运营效率分析

| 品牌/参与方 | 定位 | 核心模式 | 竞争优势 | 瓶颈与缺陷 |
| :--- | :--- | :--- | :--- | :--- |
| PLN | 国家电力公司 | 基础设施主导 | 掌控电网，资源绝对优势 | 响应迟缓，缺乏精细化商业运营，数字化能力较弱 |
| Starvo/Voltron | 私立充电运营商（CPO） | 纯进口/简易组装 | 响应速度快，品牌建立较早 | 本地含量（TKDN）低，无法参与政府补贴，准入优先级较低 |
| 本项目 | 本地集成商 | CKD深度本地化 | TKDN >= 40% -> 政策快车道 | 初始组装线建设存在一定周期 |

**运营效率指标：**
竞品平均设备可用率（Uptime）仅在60%-75%之间。本项目旨在通过OCPP 2.0协议与本地“中心-辐射”（Hub-and-Spoke）运维体系，将可用率提升至95%以上。

## 二、 准入政策与合规量化

### （一） 本地含量（TKDN）强制门槛
1. **核心阈值**
   TKDN >= 40%。
2. **政策影响分析**
   （1） TKDN < 40%（纯进口）：审批周期冗长，PLN Mobile系统优先级低，且无法参与国企（BUMN）采购。
   （2） TKDN >= 40%（本地化）：可获得SPKLU运营许可快车道，优先分配PLN核心资源站点。

### （二） 关键法规参考
1. **Perpres 79/2023**
   调整TKDN计算逻辑，将本地化率由“加分项”转变为“准入门槛”。
2. **MEMR 11/2024**
   监管权移交，SPKLU审批流程直接与本地化率挂钩。

## 三、 成本与财务基准

### （一） PLN电力成本基准
1. **零售电价 (P_retail)**：Rp 2,466/kWh（基于PLN SPKLU基准）。
2. **工业采购电价 (C_energy)**：约 Rp 1,500/kWh。
3. **单kWh毛利**：约 Rp 966。

### （二） 硬件成本CKD替代分析（以单台120kW充电桩为例）

| 部件 | 纯进口成本 (USD) | CKD本地化成本 (USD) | 成本降幅 |
| :--- | :--- | :--- | :--- |
| 整体机柜及支架 | 2,500 | 1,800 | 28% |
| 内部线缆及辅材 | 800 | 500 | 37% |
| 组装与测试费 | 1,200 | 600 | 50% |
| 合计 | 4,500 | 2,900 | 35% |

## 四、 资源协同量化
依托既有木工机械B端资源网络的预期增益如下：
1. **场址租赁成本**：预计较市场价降低30% - 50%。
2. **单站资本支出（CAPEX）**：综合降低15% - 20%。
3. **获客成本（CAC）**：利用既有信任链，成本接近于零。

---
**结论：**
量化数据表明，印尼市场正处于“需求爆发 -> 供应缺口 -> 政策强管制”的交汇点。通过实现 TKDN >= 40% 的本地化集成，并结合既有B端资源的成本对冲，本项目具备极强的财务稳健性与准入优势。
"""

manual_text = """# 印尼 EV-Charge Pro 运营执行手册

**版本：** 1.0（最终版）
**状态：** 最终同步版
**旨在：** 确保 1,000 万美元投资的战术级精准落地与标准化执行。

---

## 一、 组织架构与法律准入标准操作程序（SOP）

### （一） 法律实体设立：PT PMA 与资源协调机制
在印尼，法律实体的快速落地依赖于高效的本地协调员（Fixer）。

**设立具体流程：**
1. **协调员筛选**
   通过顶级律所（如 SSEK）或专业 PMA 咨询公司进行筛选，签署包含明确关键绩效指标（KPI，如“30日内获得NIB”）的协议。
2. **PT PMA 注册**
   拟定成立章程（Akta Pendirian）-> 注入注册资本（需符合 BKPM 最新要求，通常为 100 亿印尼盾）-> 注册 OSS 系统。
3. **获取企业登记号（NIB）**
   准确选择 KBLI 行业代码（电力安装与 EV 充电业务）-> 关联税务登记号（NPWP）。

### （二） PLN 电力升级与电网接入流程
高功率快充站须通过印尼国家电力公司（PLN）办理正式的增容申请。

| 阶段 | 动作 | 执行细节 | 关键交付物/款项 |
| :--- | :--- | :--- | :--- |
| 申请阶段 | 提交增容申请 | 通过 PLN Mobile 或线下分支机构提交申请。 | 支付接电费（Biaya Penyambungan）。 |
| 勘察阶段 | 技术实地调研 | PLN 工程师评估最近变压器位置及电缆距离。 | 确认变压器容量是否满足需求。 |
| 实施阶段 | 电缆铺设施工 | 从变压器至低压配电柜（LVMDP）的线路施工。 | 必须由 PLN 认证的承包商执行。 |
| 认证阶段 | 获取运行合格证 | 由独立检验机构验证安全与接地标准。 | 获得 SLO（Sertifikat Laik Operasi）。 |
| 激活阶段 | 挂表通电 | PLN 安装电表并正式激活线路。 | 签署完工报告（Berita Acara）。 |

## 二、 产品采购与 TKDN 本地化战略

### （一） 硬件层级配置选择
1. **旗舰站点**：采用欧美或中国一线品牌，旨在确保高可靠性与长周期寿命（5-10年）。
2. **规模站点**：采用本项目 CKD 本地化定制款，旨在实现成本最优并符合政策要求。

### （二） TKDN >= 40% 落地执行路径
**完全散件组装（CKD）执行流程：**
1. **核心部件进口**：从中国进口电源模组（Power Module）和主控板（PCB）。
2. **本地部件采购**：在印尼本地采购以下部件：
   （1） 外部机柜及壳体：于 Bekasi 或 Tangerang 钣金厂定制。
   （2） 电缆与接线端子：采购印尼本地工业标准件。
   （3） 包装与说明书：由本地印刷厂承接。
3. **本地组装集成**：模组安装 -> 内部布线 -> 固件刷写 -> QA 压力测试。
4. **认证申请流程**：提交本地采购发票与工时证明 -> 申请工业部（Kemenperin）审计 -> 获取 TKDN >= 40% 证书。

## 三、 物理安装技术指南

### （一） 场址获取与租赁（Sewa）策略
**谈判机制：**
1. **固定租金模式**：适用于高流量购物中心或加油站，每月支付固定费用。
2. **营收分成模式**：适用于新开发区域，将一定比例的充电收入支付给地主。
3. **混合模式（推荐）**：采用低底租与营收分成相结合，旨在降低初期投资风险。
4. **关键条款**：合同中必须明确约定“24小时维护进入权”。

### （二） 土建工程要求
1. **混凝土基座**
   （1） 尺寸：厚度 >= 20cm x 20cm。
   （2） 强度：至少达到 K-250 混凝土强度等级。
   （3） 养护：安装前须养护 3-7 天。
2. **接地系统**
   （1） 深度：接地极必须触及地下水位（通常为 3-6 米）。
   （2） 目标电阻：<= 5 Ohm。
3. **防水处理**
   （1） 基座高度须高于地面 10cm，以防止雨季积水。
   （2） 底座与充电桩接缝处须使用专业防水密封胶。

### （三） 电气工程标准
1. **管路安装**
   （1） 所有暴露电缆必须使用镀锌钢管（EMT）。
   （2） 地下管路埋深 30cm，下方铺设沙床以防止挤压。
2. **电缆标准**
   （1） 地下走线必须使用 NYY 或 NYFGbY（铠装电缆）。
   （2） 线径计算：基于 I = P / (V * cos phi)，并预留 20% 电压降缓冲。
3. **LVMDP 配电柜接线**
   （1） 链路顺序：主断路器 MCCB -> 浪涌保护 SPD -> RCD Type B 漏电保护 -> 分路断路器。
   （2） 标识要求：每根电缆必须使用热缩管标签（如 "EV-01-POWER"）。

## 四、 销售与运营闭环管理

### （一） 经销商激励机制（4S 店重点）

| 角色 | 单桩佣金 | 触发条件 | 支付方式 |
| :--- | :--- | :--- | :--- |
| 销售顾问 | IDR 500k - 1.5M | 安装完成并激活 | 次月随工资发放 |
| 店长 | IDR 1M - 2M | 达成季度指标 | 季度奖金 |
| 经销商负责人 | 年营收 X% | 达成年度销量目标 | 年终返利 |

### （二） 运维标准操作程序（Hub-and-Spoke 模式）
1. **区域中心（Hub）**：负责存放高价值备件（如电源模组、显示屏）。
2. **本地技术员（Spoke）**：派遣本地电工负责基础重启与电缆更换。

**24小时响应触发流程：**
1. **告警触发**：CMS 报告“离线”或用户反馈“充电失败”。
2. **初步分诊（<= 2h）**：通过云端管理系统尝试远程重启。
3. **现场派遣（<= 6h）**：若远程重启失败，派遣 Spoke 技术员现场检查。
4. **故障解决（<= 24h）**：完成部件更换或现场电源复位。

## 五、 风险控制红线

| 风险项 | 危险区域（红线） | 对冲策略 |
| :--- | :--- | :--- |
| 法律合规 | 与当地官员或地主达成口头协议。 | 铁律：无正式公司抬头 MoU 或合同，绝不采取行动。 |
| PLN 进度 | 依赖 PLN 经理的口头承诺。 | 铁律：每周举行一次正式进度会议，记录所有书面请求。 |
| 资金支付 | 向承包商预付 100% 费用。 | 铁律：采取 30% 定金 -> 40% 安装 -> 30% SLO 激活后支付。 |
| 技术安全 | 使用非 Type B 的 RCD（存在安全隐患）。 | 铁律：通电前必须对所有配电盘进行强制审计并留存影像资料。 |
| TKDN 审计 | 向工业部虚报本地化含量。 | 铁律：维护完整的 BOM 发票链，确保所有零部件可追溯。 |
"""

process_file(report_text, "/Users/openclaw/.openclaw/workspace-shrimp-commander/印尼EV充电市场量化研究报告_最终版.docx")
process_file(manual_text, "/Users/openclaw/.openclaw/workspace-shrimp-commander/印尼EV充电桩项目实战落地执行手册_最终版.docx")
