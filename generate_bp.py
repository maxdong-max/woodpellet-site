from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_bp():
    doc = Document()

    # Title Page
    title = doc.add_heading('Indonesia EV-Charge Pro: Charging Pile Distribution and Full Life Cycle Management Plan (BP)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 2)
    subtitle = doc.add_paragraph('Investment-Grade Strategic Business Plan')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 5)
    date_info = doc.add_paragraph('Date: April 17, 2026\nClassification: Confidential / Investment Grade')
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "Indonesia EV-Charge Pro aims to capitalize on the rapid electrification of the Indonesian automotive market "
        "by implementing a 'Light Asset' business model. Rather than owning the infrastructure, the company focuses on "
        "High-Value Product Sales, professional installation services, and comprehensive full-life cycle maintenance. "
        "This strategy ensures rapid scalability, minimal capital exposure, and high ROI through a specialized distribution network."
    )

    # Chapter 1: Market Analysis
    doc.add_heading('Chapter 1: Market Analysis', level=1)
    doc.add_heading('1.1 Market Growth & Dynamics', level=2)
    doc.add_paragraph(
        "The Indonesian EV charging infrastructure market is experiencing exponential growth. Current data indicates "
        "an SPKLU (Stasiun Pengisian Kendaraan Listrik Umum) growth rate of 157%, with a projected CAGR of 61.4% over the next five years. "
        "This surge is driven by the aggressive entry of global players such as BYD and VinFast, which are increasing EV "
        "adoption rates among the middle and upper-class demographics in urban centers."
    )
    
    doc.add_heading('1.2 Market Segmentation', level=2)
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Segment'
    hdr_cells[1].text = 'Target User'
    hdr_cells[2].text = 'Growth Potential'
    
    segments = [
        ('Residential', 'Private Homeowners / Luxury Apartments', 'High (Home-S focus)'),
        ('Commercial', 'Malls, Hotels, Office Buildings', 'Very High (Business-P focus)'),
        ('Public/Fleet', 'Taxi companies, Logistics hubs', 'Steady / Policy Driven')
    ]
    for seg, user, growth in segments:
        row_cells = table1.add_row().cells
        row_cells[0].text = seg
        row_cells[1].text = user
        row_cells[2].text = growth

    # Chapter 2: Product Strategy
    doc.add_heading('Chapter 2: Product Portfolio & Margin Analysis', level=1)
    doc.add_paragraph("The product lineup is divided into two core series, optimized for the 'Light Asset' model.")
    
    doc.add_heading('2.1 Home-S Series (Residential)', level=2)
    doc.add_paragraph("Designed for domestic use, emphasizing safety, aesthetics, and ease of installation.")
    # Home-S Specs
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Feature'
    hdr2[1].text = 'Specification'
    
    home_specs = [
        ('Power Output', '7kW / 11kW'),
        ('Installation Time', '<<  4 Hours'),
        ('Target Price', 'RMB 3,500 - 5,000'),
        ('Expected Margin', '35% - 45%')
    ]
    for feat, spec in home_specs:
        row = table2.add_row().cells
        row[0].text = feat
        row[1].text = spec

    doc.add_heading('2.2 Business-P Series (Commercial)', level=2)
    doc.add_paragraph("High-capacity chargers for commercial environments, featuring integrated billing and remote management.")
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = 'Feature'
    hdr3[1].text = 'Specification'
    
    biz_specs = [
        ('Power Output', '22kW / 60kW / 120kW'),
        ('Management', 'Cloud-based OCPP 1.6/2.0'),
        ('Target Price', 'RMB 15,000 - 80,000'),
        ('Expected Margin', '25% - 30%')
    ]
    for feat, spec in biz_specs:
        row = table3.add_row().cells
        row[0].text = feat
        row[1].text = spec

    # Chapter 3: Budget Allocation
    doc.add_heading('Chapter 3: Budget & Resource Allocation', level=1)
    doc.add_paragraph("Total Initial Investment: 10 Million RMB. Allocation follows the '4-3-2-1' model.")
    
    budget_data = [
        ('Product Procurement & Inventory', '4 Million RMB (40%)'),
        ('Marketing & Dealer Network', '3 Million RMB (30%)'),
        ('Operational Setup & Human Capital', '2 Million RMB (20%)'),
        ('Contingency & Reserve', '1 Million RMB (10%)')
    ]
    table4 = doc.add_table(rows=1, cols=2)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = 'Item'
    hdr4[1].text = 'Allocation'
    for item, alloc in budget_data:
        row = table4.add_row().cells
        row[0].text = item
        row[1].text = alloc

    doc.add_heading('3.1 Quarterly Spending Plan (Year 1)', level=2)
    q_table = doc.add_table(rows=1, cols=5)
    q_table.style = 'Table Grid'
    q_hdr = q_table.rows[0].cells
    q_hdr[0].text = 'Category'
    q_hdr[1].text = 'Q1 (RMB)'
    q_hdr[2].text = 'Q2 (RMB)'
    q_hdr[3].text = 'Q3 (RMB)'
    q_hdr[4].text = 'Q4 (RMB)'
    
    q_spending = [
        ('Procurement', '2,000,000', '1,000,000', '500,000', '500,000'),
        ('Marketing', '1,000,000', '800,000', '700,000', '500,000'),
        ('Operations', '500,000', '500,000', '500,000', '500,000'),
        ('Reserve', '250,000', '250,000', '250,000', '250,000')
    ]
    for cat, q1, q2, q3, q4 in q_spending:
        row = q_table.add_row().cells
        row[0].text = cat
        row[1].text = q1
        row[2].text = q2
        row[3].text = q3
        row[4].text = q4

    # Chapter 4: Growth Strategy
    doc.add_heading('Chapter 4: Growth & Market Expansion', level=1)
    doc.add_heading('4.1 Dealer Incentive Program', level=2)
    doc.add_paragraph(
        "To ensure rapid penetration, we will implement a tiered incentive system for distributors:\n"
        "- Tier 1 (Diamond): Volume > 500 units/yr. 10% rebate + Priority Technical Support.\n"
        "- Tier 2 (Gold): Volume 100-500 units/yr. 5% rebate + Marketing Co-op funds.\n"
        "- Tier 3 (Silver): Volume <<  100 units/yr. Standard margin + Training certification."
    )

    doc.add_heading('4.2 Jakarta B2B Acquisition Targets', level=2)
    doc.add_paragraph(
        "Focus will be on high-traffic commercial hubs in Jakarta (Sudirman, Kuningan, SCBD). "
        "Targets include:\n"
        "- Luxury Hotel Chains: Integration of Business-P chargers in valet areas.\n"
        "- Grade-A Office Towers: Partnership with building management for employee charging.\n"
        "- Premium Shopping Malls: Establishing 'Charging Hubs' as a value-add for shoppers."
    )

    # Chapter 5: Financial Projections
    doc.add_heading('Chapter 5: Financial Projections (3-Year)', level=1)
    fin_table = doc.add_table(rows=1, cols=4)
    fin_table.style = 'Table Grid'
    f_hdr = fin_table.rows[0].cells
    f_hdr[0].text = 'Metric'
    f_hdr[1].text = 'Year 1'
    f_hdr[2].text = 'Year 2'
    f_hdr[3].text = 'Year 3'
    
    fin_metrics = [
        ('Revenue (Million RMB)', '12.0', '35.0', '70.0'),
        ('OPEX (Million RMB)', '6.0', '12.0', '20.0'),
        ('EBITDA (Million RMB)', '2.0', '10.0', '25.0'),
        ('ROI (%)', '-10%', '45%', '110%')
    ]
    for metric, y1, y2, y3 in fin_metrics:
        row = fin_table.add_row().cells
        row[0].text = metric
        row[1].text = y1
        row[2].text = y2
        row[3].text = y3

    # Chapter 6: Risk Management
    doc.add_heading('Chapter 6: Risk Mitigation', level=1)
    doc.add_heading('6.1 TKDN (Local Content Requirement) Policy', level=2)
    doc.add_paragraph(
        "Risk: Indonesian government's increasing requirement for local content (TKDN) in electronic infrastructure. "
        "Mitigation: Phase 1 involves importing complete units; Phase 2 (Year 2) transitions to local assembly (SKD/CKD) "
        "in partnership with Indonesian manufacturers to meet the 25%-40% TKDN threshold."
    )

    doc.add_heading('6.2 Currency Fluctuations (IDR/RMB)', level=2)
    doc.add_paragraph(
        "Risk: Volatility of the Indonesian Rupiah against the RMB affecting import costs and margins. "
        "Mitigation: Use of forward contracts for large procurement cycles and implementing a dynamic pricing "
        "mechanism linked to a quarterly exchange rate average."
    )

    # Save the document
    save_path = '/Users/openclaw/.openclaw/workspace-shrimp-commander/Indonesia_EV_Charge_Pro_BP_Final.docx'
    doc.save(save_path)
    return save_path

if __name__ == '__main__':
    path = create_bp()
    print(f'Document saved to {path}')
